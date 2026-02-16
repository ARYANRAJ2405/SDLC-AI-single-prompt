from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from langgraph.graph import END, StateGraph

from .config import SYSTEM_USER_PROMPT
from .debug import debug_state
from .file_loaders import is_image_file, iter_input_files, load_source_text
from .llm_utils import generate_text_with_usage
from .models import BRDState

try:
    from docx import Document
except Exception:
    Document = None


def _fallback_brd() -> str:
    return """## Executive Summary
- TBC

## Business Context & Problem Statement
- TBC

## Objectives & Success Metrics
- TBC
"""


def _document_schema_prompt(base_prompt: str) -> str:
    return f"""{base_prompt}

Return the BRD as JSON only. Do not include markdown fences, prose, or explanations.
Use this exact schema:
{{
  "title": "string",
  "blocks": [
    {{
      "type": "heading",
      "level": 1,
      "text": "string"
    }},
    {{
      "type": "paragraph",
      "text": "string"
    }},
    {{
      "type": "bullet_list",
      "items": [
        {{"text": "string", "level": 0}},
        {{"text": "string", "level": 1}}
      ]
    }},
    {{
      "type": "numbered_list",
      "items": [
        {{"text": "string", "level": 0}},
        {{"text": "string", "level": 1}}
      ]
    }},
    {{
      "type": "table",
      "headers": ["col1", "col2"],
      "rows": [["r1c1", "r1c2"], ["r2c1", "r2c2"]]
    }}
  ]
}}

Rules:
- Include all BRD sections from the task.
- Prefer tables when information is matrix-like (roles, scope in/out, MoSCoW).
- Use list item level for indentation.
- Ensure valid JSON with double quotes only.
"""


def _extract_doc_spec(output: str) -> dict[str, Any] | None:
    text = output.strip()
    if not text:
        return None
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return None
        try:
            parsed = json.loads(text[start : end + 1])
        except json.JSONDecodeError:
            return None
    if not isinstance(parsed, dict):
        return None
    blocks = parsed.get("blocks")
    if not isinstance(blocks, list):
        return None
    return parsed


def _doc_spec_to_markdown(spec: dict[str, Any]) -> str:
    lines: list[str] = []
    title = str(spec.get("title", "")).strip()
    if title:
        lines.append(f"# {title}")
        lines.append("")

    for block in spec.get("blocks", []):
        if not isinstance(block, dict):
            continue
        block_type = block.get("type")
        if block_type == "heading":
            level = int(block.get("level", 2))
            level = min(max(level, 1), 6)
            text = str(block.get("text", "")).strip()
            if text:
                lines.append(f"{'#' * level} {text}")
                lines.append("")
        elif block_type == "paragraph":
            text = str(block.get("text", "")).strip()
            if text:
                lines.append(text)
                lines.append("")
        elif block_type in {"bullet_list", "numbered_list"}:
            items = block.get("items", [])
            if isinstance(items, list):
                idx = 1
                for item in items:
                    if isinstance(item, dict):
                        text = str(item.get("text", "")).strip()
                        level = int(item.get("level", 0))
                    else:
                        text = str(item).strip()
                        level = 0
                    if not text:
                        continue
                    indent = "  " * max(level, 0)
                    if block_type == "numbered_list":
                        lines.append(f"{indent}{idx}. {text}")
                        idx += 1
                    else:
                        lines.append(f"{indent}- {text}")
                lines.append("")
        elif block_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if isinstance(headers, list) and headers:
                header_cells = [str(h) for h in headers]
                lines.append("| " + " | ".join(header_cells) + " |")
                lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                if isinstance(rows, list):
                    for row in rows:
                        if isinstance(row, list):
                            cells = [str(cell) for cell in row]
                            if len(cells) < len(header_cells):
                                cells.extend([""] * (len(header_cells) - len(cells)))
                            lines.append("| " + " | ".join(cells[: len(header_cells)]) + " |")
                lines.append("")
    return "\n".join(lines).strip()


def _render_docx_from_spec(spec: dict[str, Any], output_path: str) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required to write .docx output.")

    doc = Document()
    title = str(spec.get("title", "")).strip()
    if title:
        doc.add_heading(title, level=1)

    for block in spec.get("blocks", []):
        if not isinstance(block, dict):
            continue
        block_type = block.get("type")
        if block_type == "heading":
            level = int(block.get("level", 2))
            level = min(max(level, 1), 4)
            text = str(block.get("text", "")).strip()
            if text:
                doc.add_heading(text, level=level)
        elif block_type == "paragraph":
            text = str(block.get("text", "")).strip()
            if text:
                doc.add_paragraph(text)
        elif block_type in {"bullet_list", "numbered_list"}:
            items = block.get("items", [])
            if isinstance(items, list):
                for item in items:
                    if isinstance(item, dict):
                        text = str(item.get("text", "")).strip()
                        level = max(int(item.get("level", 0)), 0)
                    else:
                        text = str(item).strip()
                        level = 0
                    if not text:
                        continue
                    if block_type == "numbered_list":
                        style = "List Number" if level <= 0 else f"List Number {min(level + 1, 3)}"
                    else:
                        style = "List Bullet" if level <= 0 else f"List Bullet {min(level + 1, 3)}"
                    try:
                        doc.add_paragraph(text, style=style)
                    except KeyError:
                        fallback = "List Number" if block_type == "numbered_list" else "List Bullet"
                        try:
                            doc.add_paragraph(text, style=fallback)
                        except KeyError:
                            doc.add_paragraph(text)
        elif block_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if isinstance(headers, list) and headers:
                column_count = len(headers)
                table = doc.add_table(rows=1, cols=column_count)
                table.style = "Table Grid"
                for col_idx, header in enumerate(headers):
                    table.rows[0].cells[col_idx].text = str(header)
                if isinstance(rows, list):
                    for row in rows:
                        if not isinstance(row, list):
                            continue
                        row_cells = table.add_row().cells
                        for col_idx in range(column_count):
                            cell_value = row[col_idx] if col_idx < len(row) else ""
                            row_cells[col_idx].text = str(cell_value)

    doc.save(output_path)


def generate_brd_node(state: BRDState) -> BRDState:
    print("[node] generate_brd")

    files = list(iter_input_files(state.inputs))
    texts = []
    images: list[str] = []
    for file_path in files:
        if is_image_file(file_path):
            images.append(str(file_path))
            continue
        try:
            text = load_source_text(file_path)
        except Exception:
            continue
        if text.strip():
            texts.append(f"[SOURCE: {file_path.name}]\n{text}")

    state.source_texts = texts
    state.image_paths = images
    prompt = SYSTEM_USER_PROMPT.format(INPUTS_TEXT="\n\n".join(texts))
    output, usage = generate_text_with_usage(
        _document_schema_prompt(prompt),
        max_tokens=100000,
        image_paths=state.image_paths,
    )

    print(
        f"[node] generate_brd chunk_tokens={usage['total_tokens']} "
        f"(prompt={usage['prompt_tokens']} completion={usage['completion_tokens']})"
    )

    spec = _extract_doc_spec(output)
    if spec is None:
        print("[warn] generate_brd invalid/empty JSON response; retrying once")
        retry_prompt = _document_schema_prompt(
            "Return valid JSON only. Ensure the response is a single JSON object.\n\n" + prompt
        )
        output, usage = generate_text_with_usage(
            retry_prompt,
            max_tokens=30000,
            image_paths=state.image_paths,
        )
        spec = _extract_doc_spec(output)
        print(
            f"[node] generate_brd retry_tokens={usage['total_tokens']} "
            f"(prompt={usage['prompt_tokens']} completion={usage['completion_tokens']})"
        )

    if spec is None:
        print("[warn] generate_brd still invalid; using fallback markdown")
        state.brd_markdown = _fallback_brd().strip()
    else:
        state.brd_markdown = _doc_spec_to_markdown(spec)

    if state.output_markdown_path:
        Path(state.output_markdown_path).write_text(state.brd_markdown, encoding="utf-8")
    if state.output_docx_path:
        if spec is None:
            if Document is None:
                raise RuntimeError("python-docx is required to write .docx output.")
            doc = Document()
            for line in state.brd_markdown.splitlines():
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.strip() == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(line)
            doc.save(state.output_docx_path)
        else:
            _render_docx_from_spec(spec, state.output_docx_path)

    debug_state("generate_brd", state)
    print(
        f"[node] generate_brd tokens_used={usage['total_tokens']} "
        f"(prompt={usage['prompt_tokens']} completion={usage['completion_tokens']})"
    )
    return state


def build_graph() -> StateGraph:
    graph = StateGraph(BRDState)
    graph.add_node("generate_brd", generate_brd_node)
    graph.set_entry_point("generate_brd")
    graph.add_edge("generate_brd", END)
    return graph
